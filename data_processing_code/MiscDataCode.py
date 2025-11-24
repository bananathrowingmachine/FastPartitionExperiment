"""
Stores types used across multiple files to organize data transfer, as well as manages recording disagreements without causing race conditions.

Made by bananathrowingmachine on Nov 23rd, 2025
"""
from collections import namedtuple
from dataclasses import dataclass, field
import numpy as np
from pathlib import Path
from docx import Document
import pandas as pd
from typing import Union

# Float64 max values = finfo(resolution=1e-15, min=-1.7976931348623157e+308, max=1.7976931348623157e+308, dtype=float64)
FullResultsDType = np.dtype([
    ('targetSum', np.uint64),
    ('newMemoCrazy', np.float64), 
    ('memoNormal', np.float64), 
    ('tabCrazy', np.float64),
    ('tabNormal', np.float64), 
    ('recurseNormal', np.float64) 
])

SpeedyResultsDType = np.dtype([
    ('targetSum', np.uint64),
    ('newMemoCrazy', np.float64), 
    ('oldMemoCrazy', np.float64), 
])

UnionDType = Union[FullResultsDType, SpeedyResultsDType]

ResultsWrapper = namedtuple('ResultsWrapper', ['IntCount', 'RawData', 'RecurseEstimate'])

@dataclass(frozen=True)
class ResultsWrapper:
    """
    A helpful wrapper with the current chunk of results and a few other helpful pieces of information.
    """
    IntCount: int
    RecurseEstimate: int
    RawData: np.ndarray = field(default_factory=lambda: np.empty(21, dtype=UnionDType))

@dataclass(frozen=True)
class DisagreeData:
    """
    A combined grouping of all the data needed to log disagreements.
    """
    AlgoOutputs: list[bool]
    IntCount: int
    TargetIndex: int
    TestNum: int
    TargetSum: int
    CurrentList: list[int]

@dataclass()
class DataProcessingInfo:
    """
    A combined grouping of all the different data processing data that goes with each algorithm, used solely in MainDataProcessor.py.
    """
    OfficialName: str
    DataFrame: pd.DataFrame
    BarColor: tuple[float, float, float]
    EdgeColor: tuple[float, float, float]

class DisagreeProcessor:
    def __init__(self, genFilesDir: Path):
        """
        Simple disagreement data processor object. Processes the data and stores it in a subdirectory of the one given to it during construction.

        :param genFilesDir: The directory for all generated files.
        """
        self.disagreeDir = genFilesDir / "solution_conflicts"
        self.disagreeDir.mkdir(parents=True, exist_ok=True)

    @classmethod
    def processBulkDisagreements(cls, genFilesDir: Path, dataInput: list[DisagreeData], idNum: int):
        """
        Processes a group of disagreements, and their associated data. 

        :param genFilesDir: The directory for all generated files.
        :param dataList: The list of each disagreement are the conditions were it arised in DisagreeData named tuple format.
        :param idNum: The disagreement number. Starts at the number given and will increment by one for each given disagreement.
        """
        processor = cls(genFilesDir)
        for disagreement in dataInput:
            processor.processDisagreement(disagreement, idNum)
            idNum += 1

    @classmethod
    def noDisagreements(cls, genFilesDir: Path):
        """
        Handles cases where there were no recorded disagreements.

        :param genFilesDir: The directory for the generated file.
        """
        directory = cls(genFilesDir)
        document = Document()
        document.add_heading("Complete Algorithms Disagreement Record")
        document.add_heading("This document has all recorded instances where the different partition algorithms disagreed on the answer for a given set. It is procedurally generated as the experiment runs.", 3)
        document.add_paragraph("There were no recorded disagreements throughout running the entire experiment. Therefore, there is nothing else here to see.")
        document.save(directory.disagreeDir / "DisagreementRecord.docx")

    def processDisagreement(self, data: DisagreeData, idNum: int):
        """
        Records any disagreement between the algorithms. Since the tested sets can get pretty big, all data besides the set is recorded on one txt file, and then it will point to the specific set that caused the disagreement.
        Records will list which algorithms disagreed, the current set integer count, the current target index and it's associated sum, the current test number, and the actual set was tested.

        :param data: All of the disagreement data, neatly packaged for use.
        :param idNum: The disagreement number.
        """
        xnor = data.AlgoOutputs
        algoNames = ["New Memoized Crazy: ", ", Memoized Normal: " if len(xnor) != 2 else " and Old Memoized Crazy: ", 
                     ", Tabulated Crazy: ", ("," if data.IntCount <= 25 else " and") + " Tabulated Normal: ", " and Recursive Normal: "]
        try:
            document = Document(self.disagreeDir / "DisagreementRecord.docx")
        except FileNotFoundError:
            document = Document()
            document.add_heading("Complete Algorithms Disagreement Record")
            document.add_heading("This document has all recorded instances where the different partition algorithms disagreed on the answer for a given set. It is procedurally generated as the experiment runs.", 3)

        document.add_paragraph().add_run(f"Disagreement number {idNum}:").bold = True
        
        paragraph = document.add_paragraph()
        resultString = ""
        for i in range(len(xnor)):
            resultString += algoNames[i] + str(xnor[i])
        
        paragraph.add_run(f"The results from each algorithm are; {resultString}.")
        paragraph.add_run().add_break()
        paragraph.add_run(f"The specific environment being tested when this disagreement occurred is shown below:").add_break()

        if 11 <= (data.TestNum % 100) <= 13:
            suffix = 'th'
        else:
            suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(data.TestNum % 10, 'th')
        testNumStr = f"{data.TestNum}{suffix}"
        paragraph.add_run(f"This was the {testNumStr} repeat test for this specific set of independent variables.").add_break()
        paragraph.add_run(f"The amount of integers per set was {data.IntCount}.").add_break()
        paragraph.add_run(f"The current target index was {data.TargetIndex} which corresponds to a target absolute sum of {data.TargetSum}.").add_break()
        paragraph.add_run(f"The specific set that was tested has a sum of {sum(data.CurrentList)}, and a absolute sum of {sum(map(abs, data.CurrentList))}. It is shown below:").add_break()
        paragraph.add_run(f"{data.CurrentList}")
    
        document.save(self.disagreeDir / "DisagreementRecord.docx")
        
